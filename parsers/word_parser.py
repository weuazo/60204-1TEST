from .parser_base import DocumentParser

class WordParser(DocumentParser):
    """Word 문서 파서"""
    
    def __init__(self):
        super().__init__()
        self.paragraphs = []
        self.tables = []
        
    def parse(self, file_path, **kwargs):
        """Word 파일 파싱"""
        self.file_path = file_path
        try:
            # Word 처리 라이브러리 동적 임포트
            try:
                import docx
            except ImportError:
                raise ImportError("Word 파싱을 위해 'python-docx' 라이브러리가 필요합니다. 'pip install python-docx' 명령으로 설치해주세요.")
            
            # 문서 열기
            doc = docx.Document(file_path)
            
            # 문단 추출
            self.paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # 테이블 추출
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text)
                    table_data.append(row_data)
                self.tables.append(table_data)
            
            # 메타데이터 설정
            self.metadata = {
                'file_type': 'word',
                'paragraphs': len(self.paragraphs),
                'tables': len(self.tables)
            }
            
            # 문서 속성 추출 시도
            try:
                core_props = doc.core_properties
                self.metadata.update({
                    'title': core_props.title,
                    'author': core_props.author,
                    'created': str(core_props.created) if core_props.created else None
                })
            except:
                pass  # 속성 추출 실패 무시
                
            # 토큰 추정
            self.estimate_tokens()
                
            return True
        except ImportError as e:
            raise e
        except Exception as e:
            print(f"Word 파싱 오류: {e}")
            return False
    
    def get_text_content(self):
        """모든 문단의 텍스트 반환"""
        text = "\n\n".join(self.paragraphs)
        
        # 테이블 내용 추가
        if self.tables:
            text += "\n\n[테이블]\n"
            for i, table in enumerate(self.tables):
                text += f"\n테이블 {i+1}:\n"
                for row in table:
                    text += " | ".join(row) + "\n"
        
        return text
    
    def get_structure(self):
        """Word 문서의 구조 정보 반환"""
        if not self.paragraphs:
            return None
            
        structure = {
            'type': 'word',
            'paragraph_count': len(self.paragraphs),
            'table_count': len(self.tables),
            'metadata': self.metadata
        }
        
        return structure
    
    def get_paragraphs(self):
        """모든 문단 반환"""
        return self.paragraphs
    
    def get_tables(self):
        """추출된 모든 테이블 반환"""
        return self.tables
